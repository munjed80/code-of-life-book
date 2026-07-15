#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
أداةُ بناء نسخة النشر النظيفة (publish/).

تُنتج هذه الأداة ثلاث صيغ نهائيّةٍ من الكتاب، خاليةٍ من الملاحظات التحريريّة
وعلامات الحواشي الخام والإشارات إلى ملفات المشروع الداخليّة:

1) publish/book-final.md       — نسخة نظيفة بصيغة ماركداون.
2) publish/book-final-print.txt — نصّ نظيف للطباعة.
3) publish/book-final-ebook.html — نسخة إلكترونيّة RTL بصفحة عنوان وفهرس.

الاستعمال:
    python3 tools/build_publish.py
"""

from __future__ import annotations

import html
import os
import re
from pathlib import Path
from typing import List, Tuple

ROOT = Path(__file__).resolve().parent.parent
CHAPTERS_DIR = ROOT / "chapters"
PUBLISH_DIR = ROOT / "publish"

BOOK_TITLE = "كود الحياة"
BOOK_SUBTITLE = "من النظام إلى الخالق، ومن الخالق إلى معنى الإنسان وفعله"
AUTHOR = "منجد محمد السيد"

FILES_ORDER: List[Tuple[str, str]] = [
    ("introduction.md", "المقدمة"),
    ("chapter-01.md", "الفصل الأول"),
    ("chapter-02.md", "الفصل الثاني"),
    ("chapter-03.md", "الفصل الثالث"),
    ("chapter-04.md", "الفصل الرابع"),
    ("chapter-05.md", "الفصل الخامس"),
    ("chapter-06.md", "الفصل السادس"),
    ("chapter-07.md", "الفصل السابع"),
    ("chapter-08.md", "الفصل الثامن"),
    ("chapter-09.md", "الفصل التاسع"),
    ("chapter-10.md", "الفصل العاشر"),
    ("chapter-11.md", "الفصل الحادي عشر"),
    ("chapter-12.md", "الفصل الثاني عشر"),
    ("appendix-a-formative-bridge.md", "الملحق أ"),
    ("appendix-b-symbolic-code.md", "الملحق ب"),
    ("conclusion.md", "خاتمة الكتاب"),
    ("scientific-references.md", "المراجع العلميّة"),
]

# أسماء الملفات الداخليّة التي يَجب ألا تَظهر في نسخة النشر
INTERNAL_FILES_PATTERN = (
    r"(?:BOOK_RULES|GLOSSARY|SCIENCE_NOTES|QURAN_EVIDENCE|"
    r"MATURE_THESIS|CRITIQUE|IBN_TAYMIYYAH_NOTES|CONCEPT_SEEDS|"
    r"DEEP_THINKING_NOTES|OUTLINE|README)\.md"
)

# عنوان قسم الملاحظات التحريريّة (يُحذف هو وما بعده في كل فصل)
EDITORIAL_HEADING_RE = re.compile(
    r"^\s*#{1,6}\s*(?:ملاحظات\s+تحريريّة|قائمة\s+الأدلة\s+في\s+هذا\s+الفصل)"
    r"[^\n]*$",
    re.MULTILINE,
)


# ---------------------------------------------------------------------------
# تنظيف نصّ المصدر قبل التحليل
# ---------------------------------------------------------------------------

def strip_editorial_section(md: str) -> str:
    """يَحذف قسم 'ملاحظات تحريريّة للمراجعة' وما بعده من الفصل."""
    m = EDITORIAL_HEADING_RE.search(md)
    if m:
        md = md[: m.start()].rstrip() + "\n"
    return md


def remove_internal_references(text: str) -> str:
    """يَحذف من النصّ كلَّ ما يُحيل إلى ملفات المشروع الداخليّة."""
    fn = INTERNAL_FILES_PATTERN

    # 0) إصلاحاتٌ نَصّيّةٌ مُسبقة لصياغاتٍ يُؤدّي حذفُ اسم الملفّ منها إلى
    #    جُملةٍ مَبتورة، فنُعيد صياغتَها قبل التنظيف العامّ.
    text = re.sub(
        r"يَلتزم\s+بِما\s+نَصّت\s+عليه\s+`?" + fn + r"`?\s*\([^()\n]*\)",
        "يَلتزم بقاعدةٍ من قواعد الكتاب",
        text,
    )

    # 1) أزواج أقواسٍ تَحوي اسم ملفٍ داخليّ
    text = re.sub(r"\s*\([^()\n]*?" + fn + r"[^()\n]*?\)", "", text)

    # 2) جُملةٌ بين شَرطتَين تَحوي اسمَ ملفٍ داخليّ
    text = re.sub(
        r"\s*[—–\-]\s*[^—–\n]*?" + fn + r"[^—–\n]*?[—–\-]\s*",
        " ",
        text,
    )

    # 3) صيغةٌ شائعة: "التزامًا بـ`XX.md`" أو "في `XX.md`، الباب N"
    rule_words = (
        r"(?:الباب|القاعدة|القواعد|الملحق|الفصل|الجزء|الفقرة|باب|قاعدة|الصفحة)"
    )
    file_then_tail = (
        r"(?:التزامًا\s+بـ?|في|بـ|على|عليه)\s*`?" + fn + r"`?"
        r"(?:\s*،\s*" + rule_words + r"[^،.\n:؛()—–]*)*"
        r"(?:\s*\([^()\n]*?" + rule_words + r"[^()\n]*?\))?"
    )
    text = re.sub(file_then_tail, "", text)

    # 4) أيّ ذكرٍ متبقٍّ لاسم ملفٍ داخليّ بمفرده
    text = re.sub(r"`?" + fn + r"`?", "", text)

    # 5) أقواسٌ بقايا تَحوي إحالةً إلى قاعدةٍ من قواعد الكتاب فقط — تُحذف
    #    لأنّها لم تَعد ذاتَ معنًى للقارئ بعد إزالة المراجع الداخليّة.
    bracket_only_rule = (
        r"\s*\(\s*" + rule_words + r"\s+[\d٠-٩]+[^()\n]{0,80}?\)"
    )
    text = re.sub(bracket_only_rule, "", text)

    return text


def remove_editorial_phrases(text: str) -> str:
    """يَحذف عباراتٍ تحريريّةً داخليّةً من المتن."""
    # «حالة التوثيق: ...» إلى نهاية الجملة
    text = re.sub(
        r"\s*\*{0,2}حالة\s+التوثيق\*{0,2}\s*:?[^\.\n]*\.",
        "",
        text,
    )
    # أقواسٌ تَحوي «يحتاج إلى توثيق» أو نحوَه
    text = re.sub(
        r"\s*\([^()\n]*?(?:يحتاج\s+إلى\s+توثيق|يَحتاج\s+إلى\s+توثيق|"
        r"يُحال\s+في\s+النسخة\s+النهائيّة|يُؤجَّل\s+إلى\s+الذيل\s+التحريري)"
        r"[^()\n]*?\)",
        "",
        text,
    )
    return text


def remove_footnote_markers(text: str) -> str:
    """يَحذف علاماتِ الحواشي الخام مثل [^causality] وتعريفاتِها."""
    # تعريفاتُ الحواشي في بداية الأسطر: [^xxx]: ...
    text = re.sub(
        r"^\s*\[\^[A-Za-z0-9_-]+\]:[^\n]*(?:\n[ \t]+[^\n]*)*\n?",
        "",
        text,
        flags=re.MULTILINE,
    )
    # علاماتُ الحواشي داخل المتن
    text = re.sub(r"\[\^[A-Za-z0-9_-]+\]", "", text)
    return text


def cleanup_whitespace_and_punct(text: str) -> str:
    """يُنظّف الفراغاتِ والترقيمَ بعد عمليّات الحذف."""
    # دمج المسافات المتعدّدة
    text = re.sub(r"[ \t]{2,}", " ", text)
    # إزالة فاصلة عربيّة أو لاتينيّة قبل علامة ترقيم نهايةٍ
    text = re.sub(r"\s*،\s*([\.:؛])", r"\1", text)
    # إزالة فاصلتَين عربيّتَين متتاليتَين
    text = re.sub(r"،\s*،", "،", text)
    # مسافة قبل فاصلة
    text = re.sub(r"\s+،", "،", text)
    # مسافة قبل نقطة
    text = re.sub(r"\s+\.", ".", text)
    # شَرطتان متتاليتان
    text = re.sub(r"[—–]\s*[—–]", "—", text)
    # أقواس فارغة
    text = re.sub(r"\(\s*\)", "", text)
    # أقواس عربيّة قَبل/بعد علامات ترقيم
    text = re.sub(r"\s+\)", ")", text)
    text = re.sub(r"\(\s+", "(", text)
    # إزالة أسطر فارغة زائدة
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def clean_chapter_source(md: str) -> str:
    """يُجري كلَّ مراحل التنظيف على نصّ فصلٍ خام."""
    md = strip_editorial_section(md)
    md = remove_footnote_markers(md)
    md = remove_internal_references(md)
    md = remove_editorial_phrases(md)
    md = cleanup_whitespace_and_punct(md)
    return md


# ---------------------------------------------------------------------------
# مُحلِّل ماركداون مبسَّط (مأخوذ من tools/build_print.py مع تَعديل طفيف)
# ---------------------------------------------------------------------------

def parse_blocks(md: str) -> List[dict]:
    blocks: List[dict] = []
    lines = md.splitlines()
    i = 0
    para_buf: List[str] = []

    def flush_para():
        if para_buf:
            text = " ".join(s.strip() for s in para_buf).strip()
            if text:
                blocks.append({"type": "para", "text": text})
            para_buf.clear()

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped == "":
            flush_para()
            i += 1
            continue

        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped):
            flush_para()
            blocks.append({"type": "hr"})
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush_para()
            blocks.append({
                "type": "heading",
                "level": len(m.group(1)),
                "text": m.group(2).strip(),
            })
            i += 1
            continue

        if stripped.startswith(">"):
            flush_para()
            quote_lines = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                q = re.sub(r"^\s*>\s?", "", lines[i])
                quote_lines.append(q.strip())
                i += 1
            seg = [s for s in quote_lines]
            label = None
            first = seg[0].strip() if seg else ""
            m_label = re.fullmatch(r"\*\*(.+?)\*\*", first)
            rest = [s for s in seg[1:] if s]
            if m_label and rest:
                label = m_label.group(1).strip()
                seg = seg[1:]
            text = " ".join(s for s in seg if s).strip()
            blocks.append({"type": "blockquote", "label": label, "text": text})
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            flush_para()
            blocks.append({
                "type": "list_item",
                "ordered": True,
                "number": int(m.group(1)),
                "text": m.group(2).strip(),
            })
            i += 1
            continue

        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            flush_para()
            blocks.append({
                "type": "list_item",
                "ordered": False,
                "text": m.group(1).strip(),
            })
            i += 1
            continue

        para_buf.append(stripped)
        i += 1

    flush_para()
    return blocks


def clean_inline(text: str) -> str:
    if not text:
        return text
    s = text
    s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
    for _ in range(3):
        s = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", s)
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
        s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", s)
        s = re.sub(r"___(.+?)___", r"\1", s)
        s = re.sub(r"__(.+?)__", r"\1", s)
        s = re.sub(r"(?<!_)_(?!_)([^_]+?)(?<!_)_(?!_)", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = s.replace("_", "").replace("*", "").replace("`", "")
    s = re.sub(r"[ \t]+", " ", s).strip()
    return s


def to_arabic_numerals(n: int) -> str:
    return str(n).translate(str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩"))


# ---------------------------------------------------------------------------
# توليد ماركداون نظيف (publish/book-final.md)
# ---------------------------------------------------------------------------

def build_clean_markdown(file_blocks: List[Tuple[str, List[dict]]]) -> str:
    out: List[str] = []
    out.append(f"# {BOOK_TITLE}")
    out.append("")
    out.append(f"## {BOOK_SUBTITLE}")
    out.append("")
    out.append(f"تأليف: {AUTHOR}")
    out.append("")
    out.append("---")
    out.append("")
    out.append("## فهرس الكتاب")
    out.append("")
    for _, label in FILES_ORDER:
        out.append(f"- {label}")
    out.append("")
    out.append("---")
    out.append("")

    for fname, _label in FILES_ORDER:
        blocks = dict(file_blocks).get(fname)
        if blocks is None:
            continue
        for b in blocks:
            t = b["type"]
            if t == "heading":
                level = min(max(b["level"], 1), 6)
                txt = b["text"].strip()
                if not txt:
                    continue
                out.append(f"{'#' * level} {txt}")
                out.append("")
            elif t == "para":
                out.append(b["text"].strip())
                out.append("")
            elif t == "blockquote":
                label = b.get("label")
                body = b["text"].strip()
                if label:
                    out.append(f"> **{label.strip()}**")
                    if body:
                        out.append(f"> {body}")
                else:
                    out.append(f"> {body}")
                out.append("")
            elif t == "hr":
                out.append("---")
                out.append("")
            elif t == "list_item":
                if b.get("ordered"):
                    out.append(f"{b['number']}. {b['text'].strip()}")
                else:
                    out.append(f"- {b['text'].strip()}")
        out.append("")

    text = "\n".join(out)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"


# ---------------------------------------------------------------------------
# توليد نصّ الطباعة (publish/book-final-print.txt)
# ---------------------------------------------------------------------------

def build_print_text(file_blocks: List[Tuple[str, List[dict]]]) -> str:
    out: List[str] = []
    out.append(BOOK_TITLE)
    out.append("")
    out.append(BOOK_SUBTITLE)
    out.append("")
    out.append(f"تأليف: {AUTHOR}")
    out.append("")
    out.append("")
    out.append("فهرس الكتاب")
    out.append("")
    for _, label in FILES_ORDER:
        out.append(label)
    out.append("")
    out.append("")

    for fname, _label in FILES_ORDER:
        blocks = dict(file_blocks).get(fname)
        if blocks is None:
            continue
        out.append("")
        out.append("")
        prev_type = None
        for idx, b in enumerate(blocks):
            t = b["type"]
            if t == "heading":
                txt = clean_inline(b["text"])
                if not txt:
                    continue
                if prev_type is not None:
                    out.append("")
                out.append(txt)
                out.append("")
            elif t == "para":
                out.append(clean_inline(b["text"]))
                out.append("")
            elif t == "blockquote":
                label = b.get("label")
                txt = clean_inline(b["text"])
                if label:
                    lbl = clean_inline(label)
                    if lbl and lbl[-1] in "؟!.:؛…،":
                        out.append(lbl)
                    else:
                        out.append(f"{lbl}:")
                    if txt:
                        out.append(txt)
                    out.append("")
                elif txt:
                    out.append(txt)
                    out.append("")
            elif t == "list_item":
                if b.get("ordered"):
                    num = to_arabic_numerals(b["number"])
                    out.append(f"{num}. {clean_inline(b['text'])}")
                else:
                    out.append(f"• {clean_inline(b['text'])}")
                next_type = blocks[idx + 1]["type"] if idx + 1 < len(blocks) else None
                if next_type != "list_item":
                    out.append("")
            elif t == "hr":
                out.append("")
            prev_type = t

    text = "\n".join(out)
    text = re.sub(r"\n{3,}", "\n\n\n", text)
    return text.strip() + "\n"


# ---------------------------------------------------------------------------
# توليد HTML (publish/book-final-ebook.html)
# ---------------------------------------------------------------------------

HTML_HEAD = """<!doctype html>
<html lang="ar" dir="rtl">
<head>
<meta charset="utf-8">
<title>{title} - {author}</title>
<meta name="author" content="{author}">
<meta name="viewport" content="width=device-width, initial-scale=1">
<style>
  @page {{ size: A5; margin: 22mm 18mm; }}
  html {{ font-size: 21px; }}
  body {{
    font-family: "Noto Naskh Arabic", "Amiri", "Scheherazade New", "Traditional Arabic", "Tahoma", "Arial", "Times New Roman", serif;
    font-weight: 450;
    font-size: 1rem;
    line-height: 2.05;
    color: #1a1a1a;
    background: #fafaf7;
    max-width: 880px;
    margin: 0 auto;
    padding: 2.5rem 2rem 4rem;
    text-align: justify;
    direction: rtl;
    -webkit-font-smoothing: antialiased;
    -moz-osx-font-smoothing: grayscale;
  }}
  h1, h2, h3, h4 {{
    font-weight: 800;
    line-height: 1.5;
    margin: 1.6em 0 0.6em;
    text-align: right;
    color: #0d0d0d;
  }}
  h1 {{ font-size: 2.1rem; }}
  h2 {{ font-size: 1.65rem; }}
  h3 {{ font-size: 1.3rem; }}
  h4 {{ font-size: 1.12rem; }}
  p {{ margin: 0 0 1em; text-indent: 1.2em; }}
  blockquote {{
    margin: 1.2em 0;
    padding: 0.6em 1em;
    border-inline-start: 3px solid #999;
    background: #f6f6f4;
    text-indent: 0;
  }}
  blockquote p {{ margin: 0 0 0.5em; text-indent: 0; }}
  blockquote p:last-child {{ margin-bottom: 0; }}
  blockquote .quote-label {{ font-weight: 800; margin-bottom: 0.2em; }}
  ol, ul {{ margin: 0 0 1em; padding-inline-start: 1.6em; }}
  li {{ margin-bottom: 0.35em; text-indent: 0; }}
  hr {{ border: 0; border-top: 1px solid #bbb; margin: 2em auto; width: 60%; }}
  .cover {{
    min-height: 92vh;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    text-align: center;
    page-break-after: always;
  }}
  .cover .title {{ font-size: 3rem; font-weight: 800; margin-bottom: 0.4em; }}
  .cover .subtitle {{ font-size: 1.25rem; color: #333; max-width: 28em; line-height: 1.7; margin-bottom: 3em; }}
  .cover .author-label {{ color: #555; margin-bottom: 0.2em; }}
  .cover .author {{ font-size: 1.4rem; font-weight: 700; }}
  .toc {{ page-break-after: always; }}
  .toc h2 {{ text-align: center; }}
  .toc ol {{ list-style: none; padding: 0; }}
  .toc li {{ margin: 0.5em 0; font-size: 1.05rem; }}
  .toc a {{ color: #111; text-decoration: none; }}
  .toc a:hover {{ text-decoration: underline; }}
  .section {{ page-break-before: always; }}
  .section:first-of-type {{ page-break-before: auto; }}
  @media print {{
    body {{ max-width: none; padding: 0; background: #fff; font-size: 12pt; }}
    a {{ color: inherit; text-decoration: none; }}
  }}
  @media (max-width: 640px) {{
    html {{ font-size: 20px; }}
    body {{
      padding: 1.25rem 1rem 2.5rem;
      line-height: 2;
    }}
    body, p, li {{ font-size: 19px; }}
    h1 {{ font-size: 1.7rem; }}
    h2 {{ font-size: 1.4rem; }}
    h3 {{ font-size: 1.2rem; }}
    h4 {{ font-size: 1.05rem; }}
  }}
</style>
</head>
<body>
"""

HTML_TAIL = "\n</body>\n</html>\n"


def slugify(name: str) -> str:
    base = os.path.splitext(name)[0]
    return re.sub(r"[^a-zA-Z0-9-]+", "-", base).strip("-").lower()


def render_html_blocks(blocks: List[dict]) -> str:
    out: List[str] = []
    i = 0
    n = len(blocks)
    while i < n:
        b = blocks[i]
        t = b["type"]
        if t == "heading":
            level = min(max(b["level"], 1), 4)
            out.append(f"<h{level}>{html.escape(clean_inline(b['text']))}</h{level}>")
            i += 1
        elif t == "para":
            out.append(f"<p>{html.escape(clean_inline(b['text']))}</p>")
            i += 1
        elif t == "blockquote":
            label = b.get("label")
            txt = html.escape(clean_inline(b["text"]))
            if label:
                lbl = html.escape(clean_inline(label))
                body = f"<p>{txt}</p>" if txt else ""
                out.append(
                    f"<blockquote><p class=\"quote-label\"><strong>{lbl}</strong></p>{body}</blockquote>"
                )
            else:
                out.append(f"<blockquote>{txt}</blockquote>")
            i += 1
        elif t == "hr":
            out.append("<hr>")
            i += 1
        elif t == "list_item":
            ordered = b.get("ordered", False)
            tag = "ol" if ordered else "ul"
            items = []
            while (
                i < n
                and blocks[i]["type"] == "list_item"
                and blocks[i].get("ordered", False) == ordered
            ):
                items.append(html.escape(clean_inline(blocks[i]["text"])))
                i += 1
            out.append(f"<{tag}>")
            for it in items:
                out.append(f"  <li>{it}</li>")
            out.append(f"</{tag}>")
        else:
            i += 1
    return "\n".join(out)


def build_html(file_blocks: List[Tuple[str, List[dict]]]) -> str:
    parts: List[str] = []
    parts.append(HTML_HEAD.format(
        title=html.escape(BOOK_TITLE),
        author=html.escape(AUTHOR),
    ))

    parts.append('<section class="cover">')
    parts.append(f'  <div class="title">{html.escape(BOOK_TITLE)}</div>')
    parts.append(f'  <div class="subtitle">{html.escape(BOOK_SUBTITLE)}</div>')
    parts.append('  <div class="author-label">تأليف</div>')
    parts.append(f'  <div class="author">{html.escape(AUTHOR)}</div>')
    parts.append('</section>')

    parts.append('<nav class="toc">')
    parts.append('  <h2>فهرس الكتاب</h2>')
    parts.append('  <ol>')
    fb = dict(file_blocks)
    for fname, label in FILES_ORDER:
        if fname in fb:
            parts.append(f'    <li><a href="#{slugify(fname)}">{html.escape(label)}</a></li>')
    parts.append('  </ol>')
    parts.append('</nav>')

    for fname, _label in FILES_ORDER:
        if fname not in fb:
            continue
        parts.append(f'<section class="section" id="{slugify(fname)}">')
        parts.append(render_html_blocks(fb[fname]))
        parts.append('</section>')

    parts.append(HTML_TAIL)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# تَوليد PDF (publish/book-print.pdf) — يستعمل WeasyPrint
# ---------------------------------------------------------------------------

def build_pdf(html_content: str, output_path: "Path") -> bool:
    """يُحوِّل HTML إلى PDF عبر WeasyPrint. يُعيد True عند النجاح."""
    try:
        from weasyprint import HTML as WP_HTML
        WP_HTML(string=html_content, base_url=str(PUBLISH_DIR)).write_pdf(str(output_path))
        return True
    except ImportError:
        print("[تنبيه] WeasyPrint غير مثبَّت؛ جرِّب: pip install weasyprint — تم تخطّي PDF.")
        return False
    except Exception as exc:
        print(f"[خطأ] فشل توليد PDF: {exc}")
        return False


# ---------------------------------------------------------------------------
# تَوليد DOCX (publish/book.docx)
# ---------------------------------------------------------------------------

_ARABIC_FONT = "Traditional Arabic"


def _add_rtl_run(para, text: str, font_size_pt: int = 12, bold: bool = False):
    """يُضيف run عربيًّا بخصائص RTL كاملة."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = para.add_run(text)
    run.font.name = _ARABIC_FONT
    run.font.size = Pt(font_size_pt)
    if bold:
        run.bold = True
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _ARABIC_FONT)
    rFonts.set(qn("w:hAnsi"), _ARABIC_FONT)
    rFonts.set(qn("w:cs"), _ARABIC_FONT)
    rPr.insert(0, rFonts)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size_pt * 2))
    rPr.append(szCs)
    rtl_elem = OxmlElement("w:rtl")
    rPr.append(rtl_elem)
    return run


def _set_para_rtl(para, align: str = "right"):
    """يُعيّن RTL واتّجاه النصّ على مستوى الفقرة."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.insert(0, bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), align)
    pPr.append(jc)


def build_docx(file_blocks: List[Tuple[str, List[dict]]]):
    """يبني مستند Word عربيًّا RTL من كتل الملفات المُحلَّلة."""
    try:
        from docx import Document
        from docx.shared import Cm
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("[تنبيه] python-docx غير مثبَّت؛ جرِّب: pip install python-docx — تم تخطّي DOCX.")
        return None

    doc = Document()

    # حجم الصفحة A5 وهوامش مناسبة للطباعة
    sec = doc.sections[0]
    sec.page_width = Cm(14.8)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # RTL على مستوى القسم
    bidi_sect = OxmlElement("w:bidi")
    sec._sectPr.insert(0, bidi_sect)

    # صفحة الغلاف
    p = doc.add_paragraph()
    _add_rtl_run(p, BOOK_TITLE, font_size_pt=26, bold=True)
    _set_para_rtl(p, "center")

    p = doc.add_paragraph()
    _add_rtl_run(p, BOOK_SUBTITLE, font_size_pt=13)
    _set_para_rtl(p, "center")

    p = doc.add_paragraph()
    _add_rtl_run(p, f"تأليف: {AUTHOR}", font_size_pt=14, bold=True)
    _set_para_rtl(p, "center")

    doc.add_page_break()

    # فهرس الكتاب
    p = doc.add_paragraph()
    _add_rtl_run(p, "فهرس الكتاب", font_size_pt=18, bold=True)
    _set_para_rtl(p, "center")

    fb = dict(file_blocks)
    for fname, label in FILES_ORDER:
        if fname in fb:
            p = doc.add_paragraph()
            _add_rtl_run(p, label, font_size_pt=12)
            _set_para_rtl(p)

    doc.add_page_break()

    # محتوى الكتاب
    for fname, _label in FILES_ORDER:
        blocks = fb.get(fname)
        if blocks is None:
            continue
        for b in blocks:
            t = b["type"]
            if t == "heading":
                txt = clean_inline(b["text"])
                if not txt:
                    continue
                size = {1: 20, 2: 16, 3: 14, 4: 12}.get(b["level"], 12)
                p = doc.add_paragraph()
                _add_rtl_run(p, txt, font_size_pt=size, bold=True)
                _set_para_rtl(p)
            elif t == "para":
                p = doc.add_paragraph()
                _add_rtl_run(p, clean_inline(b["text"]), font_size_pt=12)
                _set_para_rtl(p, "both")
            elif t == "blockquote":
                p = doc.add_paragraph()
                lbl = b.get("label")
                body = clean_inline(b["text"])
                if lbl:
                    _add_rtl_run(p, clean_inline(lbl) + ": ", font_size_pt=11, bold=True)
                if body:
                    _add_rtl_run(p, body, font_size_pt=11)
                _set_para_rtl(p)
                # حدٌّ أيمن للاقتباس (RTL)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                right = OxmlElement("w:right")
                right.set(qn("w:val"), "single")
                right.set(qn("w:sz"), "12")
                right.set(qn("w:space"), "4")
                right.set(qn("w:color"), "999999")
                pBdr.append(right)
                pPr.append(pBdr)
                ind = OxmlElement("w:ind")
                ind.set(qn("w:right"), "360")
                pPr.append(ind)
            elif t == "list_item":
                prefix = (
                    f"{to_arabic_numerals(b['number'])}. "
                    if b.get("ordered")
                    else "• "
                )
                p = doc.add_paragraph()
                _add_rtl_run(p, prefix + clean_inline(b["text"]), font_size_pt=12)
                _set_para_rtl(p)
            elif t == "hr":
                p = doc.add_paragraph()
                _add_rtl_run(p, "* * *", font_size_pt=10)
                _set_para_rtl(p, "center")

    return doc


# ---------------------------------------------------------------------------
# تَوليد EPUB (publish/book.epub)
# ---------------------------------------------------------------------------

_EPUB_CSS = """\
@charset "UTF-8";
html, body {
  direction: rtl;
  text-align: right;
  font-family: "Amiri", "Scheherazade New", "Traditional Arabic", "Tahoma", serif;
  line-height: 1.95;
  color: #111;
  background: #fff;
}
h1, h2, h3, h4 {
  font-weight: 700;
  line-height: 1.4;
  margin: 1.4em 0 0.5em;
  text-align: right;
}
h1 { font-size: 1.8em; }
h2 { font-size: 1.45em; }
h3 { font-size: 1.2em; }
h4 { font-size: 1.08em; }
p { margin: 0 0 0.9em; text-indent: 1.2em; }
blockquote {
  margin: 1em 0;
  padding: 0.5em 0.8em;
  border-right: 3px solid #999;
  background: #f6f6f4;
  text-indent: 0;
}
blockquote p { text-indent: 0; margin: 0 0 0.4em; }
.quote-label { font-weight: 700; }
ul, ol { padding-right: 1.6em; padding-left: 0; }
li { margin-bottom: 0.3em; text-indent: 0; }
hr { border: 0; border-top: 1px solid #bbb; margin: 1.5em auto; width: 50%; }
.cover { text-align: center; padding: 3em 1em; }
.cover .title { font-size: 2.2em; font-weight: 800; margin-bottom: 0.4em; }
.cover .subtitle { font-size: 1.2em; color: #333; margin-bottom: 2em; line-height: 1.7; }
.cover .author { font-size: 1.3em; font-weight: 700; }
"""


def _chapter_xhtml(blocks: List[dict]) -> str:
    """يبني XHTML لفصل EPUB واحد."""
    parts = [
        '<?xml version="1.0" encoding="utf-8"?>',
        '<!DOCTYPE html>',
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="ar" lang="ar" dir="rtl">',
        '<head><meta charset="utf-8"/>',
        '<link rel="stylesheet" type="text/css" href="../style/default.css"/>',
        '</head><body>',
    ]
    i = 0
    n = len(blocks)
    while i < n:
        b = blocks[i]
        t = b["type"]
        if t == "heading":
            level = min(max(b["level"], 1), 4)
            txt = html.escape(clean_inline(b["text"]))
            parts.append(f"<h{level}>{txt}</h{level}>")
            i += 1
        elif t == "para":
            parts.append(f"<p>{html.escape(clean_inline(b['text']))}</p>")
            i += 1
        elif t == "blockquote":
            lbl = b.get("label")
            txt = html.escape(clean_inline(b["text"]))
            if lbl:
                lbl_esc = html.escape(clean_inline(lbl))
                body = f"<p>{txt}</p>" if txt else ""
                parts.append(
                    f'<blockquote><p class="quote-label"><strong>{lbl_esc}</strong></p>'
                    f"{body}</blockquote>"
                )
            else:
                parts.append(f"<blockquote><p>{txt}</p></blockquote>")
            i += 1
        elif t == "hr":
            parts.append("<hr/>")
            i += 1
        elif t == "list_item":
            ordered = b.get("ordered", False)
            tag = "ol" if ordered else "ul"
            items = []
            while (
                i < n
                and blocks[i]["type"] == "list_item"
                and blocks[i].get("ordered", False) == ordered
            ):
                items.append(html.escape(clean_inline(blocks[i]["text"])))
                i += 1
            parts.append(f"<{tag}>")
            for it in items:
                parts.append(f"  <li>{it}</li>")
            parts.append(f"</{tag}>")
        else:
            i += 1
    parts.append("</body></html>")
    return "\n".join(parts)


def build_epub(file_blocks: List[Tuple[str, List[dict]]], output_path: "Path") -> bool:
    """يبني ملف EPUB عربيًّا RTL من كتل الملفات المُحلَّلة."""
    try:
        from ebooklib import epub
    except ImportError:
        print("[تنبيه] ebooklib غير مثبَّت؛ جرِّب: pip install ebooklib — تم تخطّي EPUB.")
        return False

    book = epub.EpubBook()
    book.set_identifier("code-of-life-book-ar")
    book.set_title(BOOK_TITLE)
    book.set_language("ar")
    book.add_author(AUTHOR)
    book.set_direction("rtl")

    # ورقة الأنماط
    css_item = epub.EpubItem(
        uid="style_default",
        file_name="style/default.css",
        media_type="text/css",
        content=_EPUB_CSS.encode("utf-8"),
    )
    book.add_item(css_item)

    # صفحة الغلاف
    cover_xhtml = (
        '<?xml version="1.0" encoding="utf-8"?><!DOCTYPE html>'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="ar" lang="ar" dir="rtl">'
        '<head><meta charset="utf-8"/>'
        '<link rel="stylesheet" type="text/css" href="../style/default.css"/></head>'
        '<body><div class="cover">'
        f'<div class="title">{html.escape(BOOK_TITLE)}</div>'
        f'<div class="subtitle">{html.escape(BOOK_SUBTITLE)}</div>'
        f'<div class="author">تأليف: {html.escape(AUTHOR)}</div>'
        '</div></body></html>'
    )
    cover_ch = epub.EpubHtml(title="الغلاف", file_name="cover.xhtml", lang="ar")
    cover_ch.content = cover_xhtml.encode("utf-8")
    cover_ch.add_link(href="../style/default.css", rel="stylesheet", type="text/css")
    book.add_item(cover_ch)

    # فصول المحتوى
    chapters = []
    fb = dict(file_blocks)
    for fname, label in FILES_ORDER:
        blocks = fb.get(fname)
        if blocks is None:
            continue
        ch = epub.EpubHtml(title=label, file_name=f"{slugify(fname)}.xhtml", lang="ar")
        ch.content = _chapter_xhtml(blocks).encode("utf-8")
        ch.add_link(href="../style/default.css", rel="stylesheet", type="text/css")
        book.add_item(ch)
        chapters.append(ch)

    # فهرس وعمود فقري
    book.toc = [epub.Link("cover.xhtml", "الغلاف", "cover")] + [
        epub.Link(ch.file_name, ch.title, slugify(ch.file_name)) for ch in chapters
    ]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = [cover_ch, "nav"] + chapters

    epub.write_epub(str(output_path), book)
    return True


# ---------------------------------------------------------------------------
# نقطة الدخول
# ---------------------------------------------------------------------------

def main() -> None:
    PUBLISH_DIR.mkdir(exist_ok=True)

    file_blocks: List[Tuple[str, List[dict]]] = []
    for fname, _ in FILES_ORDER:
        path = CHAPTERS_DIR / fname
        if not path.exists():
            print(f"[تحذير] ملفّ مفقود: {path}")
            continue
        md = path.read_text(encoding="utf-8")
        cleaned = clean_chapter_source(md)
        file_blocks.append((fname, parse_blocks(cleaned)))

    # ماركداون نظيف (مرجعيّ)
    md_out = build_clean_markdown(file_blocks)
    (PUBLISH_DIR / "book-final.md").write_text(md_out, encoding="utf-8")

    # نسخة الطباعة النصيّة
    txt_out = build_print_text(file_blocks)
    (PUBLISH_DIR / "book-print.txt").write_text(txt_out, encoding="utf-8")

    # النسخة الإلكترونيّة HTML
    html_out = build_html(file_blocks)
    (PUBLISH_DIR / "book-ebook.html").write_text(html_out, encoding="utf-8")

    # PDF
    pdf_ok = build_pdf(html_out, PUBLISH_DIR / "book-print.pdf")

    # DOCX
    doc = build_docx(file_blocks)
    docx_ok = False
    if doc is not None:
        doc.save(str(PUBLISH_DIR / "book.docx"))
        docx_ok = True

    # EPUB
    epub_ok = build_epub(file_blocks, PUBLISH_DIR / "book.epub")

    print(f"تم بناء نسخة النشر في: {PUBLISH_DIR}")
    print(f"  book-print.txt  {'✓' if True else '✗'}")
    print(f"  book-ebook.html {'✓' if True else '✗'}")
    print(f"  book-print.pdf  {'✓' if pdf_ok else '✗ (تحتاج: pip install weasyprint)'}")
    print(f"  book.docx       {'✓' if docx_ok else '✗ (تحتاج: pip install python-docx)'}")
    print(f"  book.epub       {'✓' if epub_ok else '✗ (تحتاج: pip install ebooklib)'}")


if __name__ == "__main__":
    main()
